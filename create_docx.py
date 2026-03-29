from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Customize the font size and font name for the new run
    new_run.font.size = Pt(9)
    new_run.font.name = 'Arial'
    new_run.font.color.rgb = RGBColor(0, 0, 0) # Black links
    new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._r)
    paragraph._p.append(hyperlink)

    return hyperlink

def add_horizontal_line(paragraph, color='CCCCCC', size=6, space=1):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)

# The new DARK RED theme color
THEME_COLOR = RGBColor(122, 24, 40) # Dark Red: #7A1828
LIGHT_THEME_COLOR = 'E6C0C5' # Light washed red for divider line C9C5E3 -> roughly #E6C0C5

doc = Document()
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(10)

# Name
p = doc.add_paragraph()
runner = p.add_run('Mohadeseh Arianrad')
runner.font.size = Pt(24)
runner.font.color.rgb = THEME_COLOR

# Job Title
p2 = doc.add_paragraph()
runner = p2.add_run('Data Analyst')
runner.font.size = Pt(13)
runner.font.color.rgb = THEME_COLOR
add_horizontal_line(p2, color='CCCCCC', size=4)

# Contact info
p3 = doc.add_paragraph()
p3.add_run('Phone: ').bold = True
p3.add_run('+49 15757171746  -  ')
p3.add_run('Email: ').bold = True
add_hyperlink(p3, 'Mohadeseh@mohadeseharianrad.de', 'mailto:Mohadeseh@mohadeseharianrad.de')
p3.add_run().add_break()
p3.add_run('LinkedIn: ').bold = True
add_hyperlink(p3, 'linkedin.com/in/mohadeseh-arianrad-30a826245', 'https://linkedin.com/in/mohadeseh-arianrad-30a826245')
p3.add_run('  -  ')
p3.add_run('Website: ').bold = True
add_hyperlink(p3, 'Mohadeseharianrad.de', 'https://Mohadeseharianrad.de')
p3.paragraph_format.space_after = Pt(14)

font = p3.runs[0].font
for r in p3.runs:
    r.font.size = Pt(9)
    r.font.name = 'Arial'

def add_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.color.rgb = THEME_COLOR
    r.font.bold = True
    add_horizontal_line(p, color=LIGHT_THEME_COLOR, size=12) # Thicker line
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_job(title, dates, company):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

    # We want Title on left, Dates on right -> Use Tabs
    tabs = p.paragraph_format.tab_stops
    # Add a right-aligned tab stop at the right margin (7.5 inches)
    from docx.enum.text import WD_TAB_ALIGNMENT
    tabs.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)

    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = 'Arial'
    
    r_tab = p.add_run('\t')
    
    r2 = p.add_run(dates)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(85, 85, 85)
    r2.font.name = 'Arial'
    
    p3 = doc.add_paragraph()
    r3 = p3.add_run(company)
    r3.bold = True
    r3.font.size = Pt(10)
    r3.font.name = 'Arial'
    p3.paragraph_format.space_after = Pt(4)

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    for r in p.runs:
        r.font.size = Pt(9.5)
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor(85, 85, 85)

# 1. Resume Summary
add_heading('Resume Summary')
p = doc.add_paragraph("Aspiring Data Analyst with strong skills in Python, SQL, and Power BI, with a focus on data modeling, predictive analytics, and process automation. Through academic projects and independent analytical work, I have developed experience in transforming complex operational and business data into meaningful insights that support data-driven decision-making. My work includes data cleaning, exploratory data analysis, and building interactive dashboards that help visualize trends and performance metrics. Continuously expanding knowledge of modern data tools, statistical techniques, and analytical methodologies to solve complex problems and create measurable value through data-driven insights.")
p.paragraph_format.space_after = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for r in p.runs:
    r.font.size = Pt(9.5)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(85, 85, 85)

# 2. Education
add_heading('Education')
p_edu1 = doc.add_paragraph()
p_edu1.paragraph_format.space_after = Pt(0)

tabs = p_edu1.paragraph_format.tab_stops
from docx.enum.text import WD_TAB_ALIGNMENT
tabs.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)

r_edu1 = p_edu1.add_run("Master's Degree in Financial Services Management")
r_edu1.bold = True
r_edu1.font.size = Pt(10)
r_edu1.font.name = 'Arial'
p_edu1.add_run("\t")
p_edu1.add_run("2025 - Present").font.size = Pt(9.5)
p_edu1.runs[-1].font.name = 'Arial'
p_edu1.runs[-1].font.color.rgb = RGBColor(85, 85, 85)


p2 = doc.add_paragraph("Kaiserslautern University of Applied Sciences")
p2.paragraph_format.space_after = Pt(2)
for r in p2.runs:
    r.font.size = Pt(9.5)
    r.font.name = 'Arial'

p3 = doc.add_paragraph("Specialization in Quantitative Methods in Finance, representing the intersection of statistical modeling and complex financial decision-making.")
p3.paragraph_format.space_after = Pt(8)
for r in p3.runs:
    r.font.size = Pt(9)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(85, 85, 85)

p4 = doc.add_paragraph()
p4.paragraph_format.space_after = Pt(0)
tabs = p4.paragraph_format.tab_stops
tabs.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)

r = p4.add_run("Bachelor's Degree in Industrial Engineering")
r.bold = True
r.font.size = Pt(10)
r.font.name = 'Arial'
p4.add_run("\t")
p4.add_run("2018 - 2022").font.size = Pt(9.5)
p4.runs[-1].font.name = 'Arial'
p4.runs[-1].font.color.rgb = RGBColor(85, 85, 85)


p5 = doc.add_paragraph("Bou Ali-Sina University, Iran")
p5.paragraph_format.space_after = Pt(2)
for r in p5.runs:
    r.font.size = Pt(9.5)
    r.font.name = 'Arial'
    
p6 = doc.add_paragraph("Mark: 2.1 | SolidWorks Instructor & Teaching Assistant")
p6.paragraph_format.space_after = Pt(12)
for r in p6.runs:
    r.font.size = Pt(9)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(85, 85, 85)

# 3. Working Experience
add_heading('Working Experience')
add_job('Quality Assurance Specialist', '08.2022 to 04.2025', 'Ghateh Sazane Razan Industrial Auto Manufacturing Co.')
add_bullet('Supported the quality team through three IATF 16949 certification audits, coordinating with 5+ departments to ensure daily work matched global automotive requirements')
add_bullet('Conducted internal audits covering 20+ different processes and 80+ products to ensure technical compliance across the production line')
add_bullet('Assisted in creating and updating 15+ procedures, forms, documents and checklists to keep documentation aligned with IATF standards, reducing manual reporting errors by roughly 15%')
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# 4. Projects & Research
add_heading('Projects & Research')
add_job('Python Automation Developer (Project)', '2023', 'Automation of Calculating Design Application')
add_bullet('Centralized database by moving all quality data into a Python-managed SQLite system, creating a single "source of truth" for the entire plant')
add_bullet('Developed a custom extraction tool to pull data from unstructured Word-based FMEA files, transforming static documentation into dynamic, actionable data')
add_bullet('Automated Risk Calculation application instantly recalculates RPNs as failure data is entered, instantly flagging high-risk areas')
add_bullet('Eliminated 99% of duplicate entries and reduced data entry time by 40%, moving from reactive reporting to real-time awareness')
doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_job('Data Analyst Researcher', '2020', 'Machine Learning and Data Mining in Supply Chain')
add_bullet('Cleaned and processed multi-stage supply chain datasets to remove reporting inconsistencies and outliers before modeling')
add_bullet('Developed a model to forecast demand and anticipate whiplash effects using Scikit-learn, achieving a low RMSE')
add_bullet('Suggested that using predictive models could reduce inventory distortion and safety stock requirements by 10–15%')

# 5. Technical Skills & Capabilities
add_heading('Technical Skills & Capabilities')
doc.add_paragraph("• Python (Tkinter)                 • SQL, SQLite                               • Pandas, Scikit-learn\n"
                  "• Power BI                             • Git Version Control                       • Relational Schema Design\n"
                  "• Data Analysis                     • IATF 16949 compliance", style='Normal').paragraph_format.space_after = Pt(12)
for r in doc.paragraphs[-1].runs:
    r.font.size = Pt(9.5)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(85, 85, 85)

# 6. Languages
add_heading('Languages')
doc.add_paragraph("• English: Fluent                    • German: ÖSD A2", style='Normal').paragraph_format.space_after = Pt(12)
for r in doc.paragraphs[-1].runs:
    r.font.size = Pt(9.5)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(85, 85, 85)

# 7. Certificates
add_heading('Certificates')
add_bullet('Generative AI for Data Analysts — Coursera')
add_bullet('Enhance your Data Analytics Career — Coursera')
add_bullet('Prompt Engineering Basics — Coursera')
add_bullet('SQL for Data Science — Coursera')
add_bullet('Machine Learning for Supply Chain — Coursera')

import os

base_dir = os.path.dirname(os.path.abspath(__file__))
assets_dir = os.path.join(base_dir, 'assets')
os.makedirs(assets_dir, exist_ok=True)

docx_path = os.path.join(assets_dir, 'Mohadeseh_Arianrad_CV.docx')
pdf_path = os.path.join(assets_dir, 'Mohadeseh_Arianrad_CV.pdf')

doc.save(docx_path)

try:
    from docx2pdf import convert
    import webbrowser
    convert(docx_path, pdf_path)
    print("Document Saved and Converted to PDF Successfully!")
    
    # Auto-open the generated PDF
    try:
        webbrowser.open(f"file://{pdf_path}")
    except Exception as e:
        print(f"Could not automatically open PDF: {e}")
        
except ImportError:
    print("Saved as docx. Please install docx2pdf to convert to pdf: pip install docx2pdf")
